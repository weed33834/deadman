"""对话增强 —— 围绕产品特色的对话侧能力

把"管理台能做的"下沉到对话里，并让对话支持文件解析：
  * POST /api/chat/upload   —— 对话中上传 PDF/Word/图片/TXT → 解析文本，供引用
  * POST /api/chat/command  —— 对话斜杠命令：
        /prompt list|get|set|new   —— 管理提示词（改人设/规则）
        /expert  list|new|delete   —— 管理自定义专家（Agent）
        /skill   list|enable|disable —— 管理技能（Skill）
  * POST /api/chat/kb         —— 查询知识库（供对话引用）

这样"在对话里改提示词/新增专家/加 skill"也能落地（与 /api/admin 同一套持久化）。
"""

from __future__ import annotations

import json
import logging
import os
import time
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Body, File, UploadFile
from fastapi.responses import Response

from ...errors import DeadmanHTTPException

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/chat", tags=["chat-extras"])

_MAX_UPLOAD = int(os.getenv("DEADMAN_CHAT_UPLOAD_MB", "25")) * 1024 * 1024


def _admin_dir() -> Path:
    d = Path.home() / ".deadman" / "admin"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _json_store(name: str) -> dict[str, Any]:
    p = _admin_dir() / f"{name}.json"
    if p.exists():
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                return data
        except (json.JSONDecodeError, OSError):
            pass
    return {}


def _json_save(name: str, data: dict[str, Any]) -> None:
    p = _admin_dir() / f"{name}.json"
    p.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


# =====================================================================
# 文件解析上传
# =====================================================================


@router.post("/upload")
async def chat_upload(
    file: UploadFile = File(default=None, description="PDF/Word/图片/TXT"),  # noqa: B008
) -> dict[str, Any]:
    """POST /api/chat/upload —— 对话中上传并解析文件，返回可引用文本。

    支持：.pdf .docx .txt .md .csv 及常见图片（OCR 可选）。
    """
    from ...doc_extract.extractor import DocumentExtractor

    if file is None:
        raise DeadmanHTTPException("DM-VOICE-4001", message="缺少文件")
    content = await file.read()
    if len(content) == 0:
        raise DeadmanHTTPException("DM-VOICE-4001", message="文件为空")
    if len(content) > _MAX_UPLOAD:
        raise DeadmanHTTPException(
            "DM-VOICE-4130", message=f"文件过大（上限 {_MAX_UPLOAD // 1024 // 1024}MB）"
        )
    filename = file.filename or "upload.bin"
    try:
        extractor = DocumentExtractor()
        file_type = extractor._detect_file_type(filename, content)
        text = extractor._extract_text(content, file_type)
        text = (text or "").strip()
        return {
            "ok": True,
            "file_name": filename,
            "file_type": file_type,
            "size": len(content),
            "char_count": len(text),
            "text": text[:20000],  # 截断，避免超长进上下文
            "truncated": len(text) > 20000,
            "hint": "可在对话中直接引用该文件内容；要点：帮助确认文件类型后再追问。",
        }
    except Exception as exc:
        logger.warning("chat_upload 解析失败 %s: %s", filename, exc)
        raise DeadmanHTTPException("DM-VOICE-5000", message=f"文件解析失败: {exc}") from exc


# =====================================================================
# 对话命令（管理提示词 / 专家 / skill）
# =====================================================================

_PROMPT_HELP = (
    "用法: /prompt list | /prompt get <name> | /prompt set <name> <内容> | /prompt new <name> <内容>\n"
    "内置资源: "
    + ", ".join(
        [
            "death-aftercare",
            "legal-advisor",
            "financial-analyst",
            "policy-researcher",
            "cross-border-specialist",
            "medical-guide",
        ]
    )
)
_EXPERT_HELP = "用法: /expert list | /expert new <id> <名称> <人设> | /expert delete <id>"
_SKILL_HELP = "用法: /skill list | /skill enable <name> | /skill disable <name>"


def _builtin_prompt_names() -> list[str]:
    from ...config import settings

    names: list[str] = []
    for d in (settings.agents_dir, settings.rules_dir):
        if d.exists():
            names += [f.stem for f in d.glob("*.md")]
    return names


def _cmd_prompt(tokens: list[str]) -> dict[str, Any]:
    action = tokens[0] if tokens else "help"
    if action == "help":
        return {"ok": True, "text": _PROMPT_HELP, "kind": "text"}
    if action == "list":
        custom = list(_json_store("prompts").keys())
        return {
            "ok": True,
            "kind": "list",
            "items": {"custom": custom, "builtin": _builtin_prompt_names()},
        }
    if action == "get" and len(tokens) >= 2:
        name = tokens[1]
        store = _json_store("prompts")
        content = store.get(name)
        if content is None:
            return {
                "ok": False,
                "kind": "text",
                "text": f"提示词 {name} 不存在（内置请看 /prompt list）",
            }
        return {"ok": True, "kind": "text", "text": f"【{name}】\n{content.get('content', '')}"}
    if action in ("set", "new") and len(tokens) >= 3:
        name, content = tokens[1], " ".join(tokens[2:])
        store = _json_store("prompts")
        existing = store.get(name) or {}
        existing["content"] = content
        existing.setdefault("version", 0)
        existing["version"] += 1
        existing["description"] = existing.get("description", "")
        store[name] = existing
        _json_save("prompts", store)
        return {
            "ok": True,
            "kind": "text",
            "text": f"已{'更新' if action == 'set' else '新建'}提示词 {name}（v{existing['version']}）。管理台「提示词」面板可查看。",
        }
    return {"ok": False, "kind": "text", "text": _PROMPT_HELP}


def _cmd_expert(tokens: list[str]) -> dict[str, Any]:
    action = tokens[0] if tokens else "help"
    if action == "help":
        return {"ok": True, "text": _EXPERT_HELP, "kind": "text"}
    if action == "list":
        custom = list(_json_store("agents").keys())
        builtin = [
            "death-aftercare",
            "legal-advisor",
            "financial-analyst",
            "policy-researcher",
            "cross-border-specialist",
            "medical-guide",
        ]
        return {"ok": True, "kind": "list", "items": {"custom": custom, "builtin": builtin}}
    if action == "new" and len(tokens) >= 3:
        eid, name = tokens[1], tokens[2]
        prompt = (
            " ".join(tokens[3:])
            if len(tokens) > 3
            else "你是一位专业的助手，请认真、准确、负责地回答用户问题。"
        )
        store = _json_store("agents")
        store[eid] = {
            "id": eid,
            "name": name,
            "system_prompt": prompt,
            "type": "custom",
            "temperature": 0.3,
            "max_steps": 10,
        }
        _json_save("agents", store)
        return {
            "ok": True,
            "kind": "text",
            "text": f"已新增专家 {eid}（{name}）。管理台「Agent」面板可查看。",
        }
    if action == "delete" and len(tokens) >= 2:
        eid = tokens[1]
        if eid in (
            "death-aftercare",
            "legal-advisor",
            "financial-analyst",
            "policy-researcher",
            "cross-border-specialist",
            "medical-guide",
        ):
            return {"ok": False, "kind": "text", "text": "内置专家不可删除"}
        store = _json_store("agents")
        if eid in store:
            del store[eid]
            _json_save("agents", store)
            return {"ok": True, "kind": "text", "text": f"已删除专家 {eid}"}
        return {"ok": False, "kind": "text", "text": f"专家 {eid} 不存在"}
    return {"ok": False, "kind": "text", "text": _EXPERT_HELP}


def _cmd_skill(tokens: list[str]) -> dict[str, Any]:
    action = tokens[0] if tokens else "help"
    if action == "help":
        return {"ok": True, "text": _SKILL_HELP, "kind": "text"}
    try:
        from ...config import settings
        from ...marketplace.skill_manager import get_skill_manager

        mgr = get_skill_manager(settings.skills_dir)
        skills = mgr.list_skills() if hasattr(mgr, "list_skills") else []
    except Exception as exc:
        logger.debug("skill 命令取技能失败: %s", exc)
        skills = []
    state = _skill_state()  # {name: enabled}
    if action == "list":
        names = []
        for s in skills:
            name = getattr(s, "name", None) or (s.get("name") if isinstance(s, dict) else str(s))
            state.setdefault(name, True)
            names.append(f"{name}{'' if state[name] else '（停用）'}")
        _skill_state_save(state)
        return {"ok": True, "kind": "list", "items": {"skills": names}}
    if action in ("enable", "disable"):
        name = tokens[1] if len(tokens) >= 2 else ""
        if not name:
            return {"ok": False, "kind": "text", "text": "用法: /skill enable|disable <名称>"}
        enabled = action == "enable"
        state[name] = enabled
        _skill_state_save(state)
        return {
            "ok": True,
            "kind": "text",
            "text": f"技能 {name} 已{'启用' if enabled else '停用'}（状态持久化，管理台技能面板可见）",
        }
    return {"ok": False, "kind": "text", "text": _SKILL_HELP}


def _skill_state() -> dict[str, bool]:
    p = _admin_dir() / "skills_state.json"
    if p.exists():
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                return data
        except (json.JSONDecodeError, OSError):
            pass
    return {}


def _skill_state_save(state: dict[str, bool]) -> None:
    (_admin_dir() / "skills_state.json").write_text(
        json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8"
    )


@router.post("/command")
async def chat_command(
    command: str = Body(default=None, embed=True, description="斜杠命令"),
) -> dict[str, Any]:
    """POST /api/chat/command —— 解析并执行对话斜杠命令。"""
    command = (command or "").strip()
    if not command.startswith("/"):
        return {"ok": False, "kind": "text", "text": "命令需以 / 开头（如 /prompt help）"}
    parts = command[1:].split()
    if not parts:
        return {"ok": False, "kind": "text", "text": "命令为空。可用: /prompt /expert /skill"}
    cmd, tokens = parts[0].lower(), parts[1:]
    if cmd == "help":
        return await _cmd_help()
    if cmd == "prompt":
        return _cmd_prompt(tokens)
    if cmd == "expert":
        return _cmd_expert(tokens)
    if cmd == "skill":
        return _cmd_skill(tokens)
    if cmd in ("hotline", "hotlines"):
        return _cmd_hotline(tokens)
    if cmd in ("institution", "institutions", "org"):
        return _cmd_institution(tokens)
    if cmd in ("custom", "customs", "民俗"):
        return _cmd_custom(tokens)
    if cmd in ("family", "kinship", "亲属"):
        return _cmd_family(tokens)
    if cmd in ("vault", "legacy", "资产"):
        if tokens and tokens[0] == "add":
            return await _cmd_vault_add(tokens[1:])
        return _cmd_vault(tokens)
    if cmd in ("note", "ending", "终活"):
        if tokens and tokens[0] == "set":
            return await _cmd_note_set(tokens[1:])
        return _cmd_note(tokens)
    if cmd in ("task", "schedule", "定时"):
        return await _cmd_task(tokens)
    if cmd in ("switch", "死人开关"):
        return _cmd_switch(tokens)
    if cmd in ("docs", "文档"):
        return _cmd_docs(tokens)
    if cmd in ("memorial", "悼文"):
        return await _cmd_memorial(tokens)
    return {
        "ok": False,
        "kind": "text",
        "text": f"未知命令 /{cmd}。可用: /prompt /expert /skill /hotline /institution /custom /family /vault /note /task /switch /docs /plot /image /browse",
    }


def _cmd_hotline(tokens: list[str]) -> dict[str, Any]:
    """查询官方热线：/hotline [省份] [功能] 或 /hotline 功能"""
    try:
        from ...hotlines.lookup import HotlineLookup

        lookup = HotlineLookup()
        province = tokens[0] if len(tokens) >= 1 else None
        function = tokens[1] if len(tokens) >= 2 else None
        results = lookup.lookup(province, function)
        if not results:
            return {
                "ok": True,
                "kind": "text",
                "text": f"未找到热线（省份={province or '全国'}，功能={function or '全部'}）。试试 /hotline 北京 殡葬",
            }
        md = ["**官方热线查询**\n"]
        for r in results[:12]:
            name = r.get("name") or r.get("机构") or "—"
            phone = r.get("phone") or r.get("电话") or "—"
            fn = r.get("function") or r.get("功能") or ""
            md.append(f"- **{name}**：{phone}{('（' + fn + '）') if fn else ''}")
        return {"ok": True, "kind": "text", "text": "\n".join(md)}
    except Exception as exc:
        return {"ok": False, "kind": "text", "text": f"热线查询失败: {exc}"}


def _cmd_institution(tokens: list[str]) -> dict[str, Any]:
    """查询机构：/institution [省份] [城市] [类型|关键词]"""
    try:
        from ...institutions.store import InstitutionStore

        store = InstitutionStore()
        province = tokens[0] if len(tokens) >= 1 else None
        city = tokens[1] if len(tokens) >= 2 else None
        keyword = tokens[2] if len(tokens) >= 3 else None
        results = store.search(province, city, None, keyword)
        if not results:
            return {
                "ok": True,
                "kind": "text",
                "text": f"未找到机构（省份={province or '-'}，城市={city or '-'}）。试试 /institution 北京 殡仪馆",
            }
        md = ["**机构查询**\n"]
        for r in results[:12]:
            d = r.to_dict() if hasattr(r, "to_dict") else r
            md.append(
                f"- **{d.get('name', '—')}**：{d.get('address', '')} {d.get('phone', '')}".rstrip()
            )
        return {"ok": True, "kind": "text", "text": "\n".join(md)}
    except Exception as exc:
        return {"ok": False, "kind": "text", "text": f"机构查询失败: {exc}"}


# =====================================================================
# 知识库引用
# =====================================================================


@router.post("/export")
async def chat_export(
    text: str = Body(default=None, embed=True, description="要导出的文本/Markdown"),
    format: str = Body(default="md", description="导出格式：md / docx / pdf"),
    filename: str = Body(default="导出内容"),
) -> Response:
    """POST /api/chat/export —— 把对话内容导出为 MD / Word(docx) / PDF。

    用于把 AI 生成的方案/清单/悼文等导出为正式文件。
    """

    text = text or ""
    fmt = (format or "md").lower()
    if fmt == "md":
        data = text.encode("utf-8")
        media = "text/markdown; charset=utf-8"
        name = f"{filename}.md"
    elif fmt == "docx":
        data = _build_docx(text)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        name = f"{filename}.docx"
    elif fmt == "pdf":
        data = _build_pdf(text)
        media = "application/pdf"
        name = f"{filename}.pdf"
    else:
        raise DeadmanHTTPException("DM-VALID-4001", message="格式仅支持 md / docx / pdf")
    return Response(
        content=data,
        media_type=media,
        headers={"Content-Disposition": _content_disposition(name)},
    )


def _content_disposition(name: str) -> str:
    """生成兼容中文文件名的 Content-Disposition（RFC 5987 filename*）。"""
    from urllib.parse import quote

    ascii_name = "".join(c for c in name if ord(c) < 128) or "export"
    return f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{quote(name)}"


def _build_docx(text: str) -> bytes:
    """生成 .docx 字节（python-docx，未安装则降级为最小 XML zip）。"""
    import io

    try:
        from docx import Document

        doc = Document()
        for para in text.split("\n"):
            line = para.strip()
            if not line:
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith(("- ", "* ")):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        # 降级：最小合法 .docx（zip 含 document.xml）
        import zipfile

        xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body>"
            + "".join(f"<w:p><w:r><w:t>{_xml_esc(p)}</w:t></w:r></w:p>" for p in text.split("\n"))
            + "</w:body></w:document>"
        )
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            z.writestr(
                "[Content_Types].xml",
                '<?xml version="1.0" encoding="UTF-8"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>',
            )
            z.writestr(
                "_rels/.rels",
                '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>',
            )
            z.writestr("word/document.xml", xml)
        return buf.getvalue()


def _build_pdf(text: str) -> bytes:
    """生成 PDF 字节（reportlab + 内置 CJK 字体支持中文）。"""
    import io

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        styles = getSampleStyleSheet()
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        story = []
        for para in text.split("\n"):
            line = para.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
            style = styles["BodyText"]
            if line.startswith("### "):
                line, style = line[4:], styles["Heading3"]
            elif line.startswith("## "):
                line, style = line[3:], styles["Heading2"]
            elif line.startswith("# "):
                line, style = line[2:], styles["Heading1"]
            story.append(Paragraph(_xml_esc(line), style))
            story.append(Spacer(1, 4))
        doc.build(story)
        return buf.getvalue()
    except Exception:
        # 降级：极简 PDF（仅 ASCII 文本；中文环境建议 reportlab 正常路径）
        lines = [_pdf_esc(p) for p in text.split("\n") if p.strip()][:30]
        body = "".join(
            f"BT /F1 11 Tf 50 {550 - i * 16} Td ({ln}) Tj ET\n" for i, ln in enumerate(lines)
        )
        pdf = (
            "%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            "2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
            "3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 595 842]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
            f"4 0 obj<</Length {len(body)}>>stream\n{body}endstream\nendobj\n"
            "5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
            "trailer<</Root 1 0 R>>\n%%EOF"
        ).encode("latin-1")
        return pdf


def _xml_esc(s: str) -> str:
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_esc(s: str) -> str:
    return str(s).replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


@router.post("/kb")
async def chat_kb(
    query: str = Body(default=None, embed=True, description="查询"),
    country: str = Body(default="CN"),
    region: str = Body(default=""),
    top_k: int = Body(default=3, ge=1, le=10),
) -> dict[str, Any]:
    """POST /api/chat/kb —— 检索知识库（供对话引用政策信息）"""
    from ...mcp_server.server import mcp

    try:
        result = await mcp.call_tool(
            "query_knowledge", {"country": country, "topic": query or "", "region": region or None}
        )
        # 提取引用来源（full_file），供前端展示"引用来源"
        sources: list[str] = []
        data = result.get("data") if isinstance(result, dict) else None
        if isinstance(data, dict) and data.get("full_file"):
            sources.append(data["full_file"])
        result["sources"] = sources
        return {"ok": True, "result": result, "sources": sources}
    except Exception as exc:
        raise DeadmanHTTPException("DM-TEXT-4040", message=f"知识库查询失败: {exc}") from exc


@router.post("/plot")
async def chat_plot(
    code: str = Body(default=None, embed=True, description="Python 绘图代码"),
    timeout: int = Body(default=30),
) -> dict[str, Any]:
    """POST /api/chat/plot —— 运行绘图代码，返回生成的图表（base64）。

    供对话画图：用户写一段 matplotlib 代码 → 沙箱执行 → 返回图片。
    """
    from ...mcp_server.server import mcp

    result = await mcp.call_tool("execute_code", {"code": code or "", "timeout": timeout})
    img = result.get("image_base64")
    if not img:
        return {
            "ok": False,
            "message": "代码未生成图片（请确认使用 matplotlib 并调用 plt.show()）",
            "stdout": result.get("stdout", ""),
            "stderr": result.get("stderr", ""),
        }
    return {
        "ok": True,
        "image_base64": img,
        "image_mime": "image/png",
        "stdout": result.get("stdout", ""),
    }


@router.post("/image")
async def chat_image(
    prompt: str = Body(default=None, embed=True, description="图像描述"),
    style: str = Body(
        default="memorial_card", description="风格：memorial_card/obituary/portrait/condolence_card"
    ),
) -> dict[str, Any]:
    """POST /api/chat/image —— 生成图片（AI 图像生成），返回 base64。

    供对话 /image 命令使用（围绕产品特色：纪念卡/讣告/肖像等风格）。
    """
    import asyncio
    import base64

    from ...multimodal.image_gen import ImageStyle, get_image_generator

    if not prompt:
        raise DeadmanHTTPException("DM-VALID-4002", message="prompt 必填")
    try:
        style_enum = ImageStyle(style)
    except ValueError:
        style_enum = ImageStyle.MEMORIAL_CARD
    gen = get_image_generator()
    if not gen.is_enabled():
        raise DeadmanHTTPException(
            "DM-VOICE-5030", message="图像生成未启用（DEADMAN_MULTIMODAL_ENABLED=0）"
        )
    img_bytes = await asyncio.to_thread(gen.generate, prompt, style_enum)
    if not img_bytes:
        raise DeadmanHTTPException("DM-VOICE-5000", message="图像生成失败（无输出）")
    return {
        "ok": True,
        "image_base64": base64.b64encode(img_bytes).decode(),
        "image_mime": "image/png",
        "prompt": prompt,
        "style": style_enum.value,
    }


@router.post("/browse")
async def chat_browse(
    url: str = Body(default=None, embed=True, description="要浏览的网页 URL"),
) -> dict[str, Any]:
    """POST /api/chat/browse —— 浏览器自动化（抓取网页并提取可读文本/摘要）

    供对话 /browse 命令：让 Agent 抓取并总结网页内容。
    """
    if not url or not url.startswith(("http://", "https://")):
        raise DeadmanHTTPException("DM-VALID-4001", message="url 需以 http(s):// 开头")
    import httpx

    try:
        async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
            resp = await client.get(url, headers={"User-Agent": "Mozilla/5.0 (deadman-agent)"})
            resp.raise_for_status()
    except Exception as exc:
        raise DeadmanHTTPException("DM-VOICE-5000", message=f"网页抓取失败: {exc}") from exc
    # 提取可读文本
    text = _extract_readable(resp.content)
    return {
        "ok": True,
        "url": url,
        "status": resp.status_code,
        "title": _extract_title(resp.content),
        "text": text[:8000],
        "truncated": len(text) > 8000,
        "hint": "可让 Agent 结合该网页内容继续回答；或在对话发 /browse <url> 触发。",
    }


def _extract_readable(html: bytes) -> str:
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        return "\n".join(lines)
    except Exception:
        return html.decode("utf-8", errors="replace")[:8000]


def _extract_title(html: bytes) -> str:
    try:
        from bs4 import BeautifulSoup

        return (BeautifulSoup(html, "html.parser").title.string or "").strip()
    except Exception:
        return ""


# =====================================================================
# 民俗规则 / 亲属图谱 对话命令
# =====================================================================

_CUSTOM_HELP = (
    "用法: /custom list | /custom get <地区|关键词> | /custom import <预置id> | /custom presets"
)
_FAMILY_HELP = "用法: /family list | /family add <姓名> [男|女] | /family rel <甲> <乙> <配偶|父|母|子|女|兄弟>"


def _cmd_custom(tokens: list[str]) -> dict[str, Any]:
    action = tokens[0] if tokens else "help"
    try:
        from ...web.routes.customs import _store as _cstore
    except Exception:
        return {"ok": False, "kind": "text", "text": "民俗模块不可用"}
    if action == "help":
        return {"ok": True, "kind": "text", "text": _CUSTOM_HELP}
    if action == "presets":
        from ...web.routes.customs import _PRESETS

        return {
            "ok": True,
            "kind": "text",
            "text": "可导入预置：\n" + "\n".join(f"- {p['id']}（{p['title']}）" for p in _PRESETS),
        }
    if action == "import" and len(tokens) >= 2:
        return {
            "ok": True,
            "kind": "text",
            "text": f"请在管理台「民俗规则」面板导入预置 {tokens[1]}，或稍后用 /custom get 查询已导入内容。",
        }
    if action == "get":
        q = " ".join(tokens[1:]) if len(tokens) > 1 else ""
        items = [dict(v, id=k) for k, v in _cstore().items() if isinstance(v, dict)]
        if q:
            items = [
                i
                for i in items
                if q in i.get("title", "")
                or q in i.get("region", "")
                or any(r.get("title", "") in q for r in i.get("rules", []))
            ]
        if not items:
            return {
                "ok": True,
                "kind": "text",
                "text": f"未找到相关民俗（关键词={q or '全部'}）。试试 /custom list 或到管理台导入预置。",
            }
        md = ["**民俗规则查询**\n"]
        for i in items[:5]:
            md.append(f"### {i['title']}（{i.get('region', '')}）")
            if i.get("process"):
                md.append("**流程**：" + " → ".join(i["process"][:8]))
            for r in i.get("rules", [])[:6]:
                md.append(f"- **{r.get('title', '')}**：{r.get('detail', '')}")
            if i.get("weekly_observances"):
                md.append(
                    "**烧七/祭奠**："
                    + "、".join(
                        f"{w.get('day', '')}({w.get('note', '')})"
                        for w in i["weekly_observances"][:7]
                    )
                )
        return {"ok": True, "kind": "text", "text": "\n".join(md)}
    if action == "list":
        items = [dict(v, id=k) for k, v in _cstore().items() if isinstance(v, dict)]
        if not items:
            return {
                "ok": True,
                "kind": "text",
                "text": "暂无已导入民俗。请在管理台「民俗规则」导入预置，或输入 /custom presets 查看可导入项。",
            }
        return {
            "ok": True,
            "kind": "text",
            "text": "已配置民俗：\n"
            + "\n".join(f"- **{i['title']}**（{i.get('region', '')}）" for i in items[:10]),
        }
    return {"ok": False, "kind": "text", "text": _CUSTOM_HELP}


def _cmd_family(tokens: list[str]) -> dict[str, Any]:
    action = tokens[0] if tokens else "help"
    try:
        from ...web.routes.kinship import _load as _kload
        from ...web.routes.kinship import _save as _ksave
    except Exception:
        return {"ok": False, "kind": "text", "text": "亲属图谱模块不可用"}
    if action == "help":
        return {"ok": True, "kind": "text", "text": _FAMILY_HELP}
    if action == "add" and len(tokens) >= 2:
        name = tokens[1]
        gender = (
            "male"
            if ("男" in name or (len(tokens) > 2 and "男" in tokens[2]))
            else (
                "female" if ("女" in name or (len(tokens) > 2 and "女" in tokens[2])) else "unknown"
            )
        )
        data = _kload()
        data["members"].append(
            {
                "id": f"m-{int(time.time() * 1000) % 10**8}",
                "name": name,
                "gender": gender,
                "note": "",
                "birth": "",
            }
        )
        _ksave(data)
        return {
            "ok": True,
            "kind": "text",
            "text": f"已加入亲属：{name}（共 {len(data['members'])} 人）。可在「亲属图谱」页查看可视化。",
        }
    if action == "list":
        data = _kload()
        members = data.get("members", [])
        if not members:
            return {"ok": True, "kind": "text", "text": "暂无亲属成员。用 /family add 姓名 添加。"}
        return {
            "ok": True,
            "kind": "text",
            "text": "亲属成员：\n"
            + "\n".join(
                f"- **{m.get('name', '')}**（{'男' if m.get('gender') == 'male' else '女' if m.get('gender') == 'female' else '未知'}）"
                for m in members[:20]
            ),
        }
    return {"ok": False, "kind": "text", "text": _FAMILY_HELP}


# =====================================================================
# 对话调动剩余功能：/vault /note /task /switch /docs
# =====================================================================

_USER = "default"  # dev/anonymous 用 default 用户（对话命令无认证上下文）


def _cmd_vault(tokens: list[str]) -> dict[str, Any]:
    """数字遗产保险库：/vault list"""
    try:
        from ...digital_legacy import DigitalLegacyStore
    except Exception:
        try:
            from ...vault.store import VaultStore
        except Exception:
            return {"ok": False, "kind": "text", "text": "保险库模块不可用"}
        store = VaultStore()
        try:
            items = store.list_items(_USER, _USER)
            md = ["**数字遗产保险库**\n"]
            for it in items[:12]:
                md.append(f"- **{it.get('title', '?')}**（{it.get('type', '')}）")
            return {
                "ok": True,
                "kind": "text",
                "text": "\n".join(md) if len(items) else "保险库暂无条目。可在「保险库」页添加。",
            }
        except Exception as exc:
            return {"ok": True, "kind": "text", "text": f"保险库：{exc}"}
    store = DigitalLegacyStore(user_id=_USER)
    try:
        s = store.summary()
        md = ["**数字遗产清单**\n"]
        md.append(f"- 资产：{s.get('total_assets', 0)} 项，未指派 {s.get('unassigned', 0)} 项")
        return {"ok": True, "kind": "text", "text": "\n".join(md)}
    except Exception as exc:
        return {"ok": False, "kind": "text", "text": f"保险库读取失败: {exc}"}


def _cmd_note(tokens: list[str]) -> dict[str, Any]:
    """终活笔记：/note list"""
    try:
        from ...ending_note.store import EndingNoteStore

        note = EndingNoteStore().load(_USER)
        if note is None:
            return {
                "ok": True,
                "kind": "text",
                "text": "尚无终活笔记。可在「终活笔记」页开始填写。",
            }
        sections = (
            note.sections
            if hasattr(note, "sections")
            else (note.get("sections", {}) if isinstance(note, dict) else {})
        )
        done = len([s for s in sections.values() if s]) if isinstance(sections, dict) else 0
        total = len(sections) if isinstance(sections, dict) else 0
        return {
            "ok": True,
            "kind": "text",
            "text": f"终活笔记：已完成 {done}/{total} 章节。可在「终活笔记」页继续填写。",
        }
    except Exception as exc:
        return {"ok": False, "kind": "text", "text": f"终活笔记读取失败: {exc}"}


async def _cmd_task(tokens: list[str]) -> dict[str, Any]:
    """定时任务：/task list | /task add <cron> <内容>"""
    action = tokens[0] if tokens else "list"
    try:
        from ...cron.scheduler import CronScheduler

        sched = CronScheduler()
    except Exception as exc:
        return {"ok": False, "kind": "text", "text": f"定时任务模块不可用: {exc}"}
    if action == "add":
        # cron 为 5 个字段（可能被空格拆开），重组
        cron = (
            " ".join(tokens[1:6]) if len(tokens) >= 6 else (tokens[1] if len(tokens) >= 2 else "")
        )
        content = (
            " ".join(tokens[6:])
            if len(tokens) >= 7
            else (" ".join(tokens[2:]) if len(tokens) > 2 else "")
        )
        if not cron or len(cron.split()) != 5:
            return {
                "ok": False,
                "kind": "text",
                "text": "用法: /task add <cron 5字段> <内容>，如 /task add 0 9 * * * 提醒办理",
            }
        try:
            res = await sched.propose_job(_USER, cron, content)
            return {
                "ok": True,
                "kind": "text",
                "text": f"已提议定时任务（需在管理台确认）：{res.get('message', '')}",
            }
        except Exception as exc:
            return {"ok": False, "kind": "text", "text": f"提议失败: {exc}"}
    jobs = sched.list_jobs(_USER)
    if not jobs:
        return {
            "ok": True,
            "kind": "text",
            "text": '暂无定时任务。用法: /task add "0 9 * * *" 提醒内容',
        }
    md = ["**定时任务**\n"]
    for j in jobs[:10]:
        d = j.to_dict() if hasattr(j, "to_dict") else j
        md.append(
            f"- {d.get('schedule', '')} · {d.get('content', '')} · {'已确认' if d.get('confirmed') else '待确认'}"
        )
    return {"ok": True, "kind": "text", "text": "\n".join(md)}


def _cmd_switch(tokens: list[str]) -> dict[str, Any]:
    """Dead Man Switch：/switch status"""
    try:
        from ...deadman_switch import store as sw_store

        st = sw_store.SwitchStore()
        try:
            status = st.get_status(_USER) if hasattr(st, "get_status") else None
        except Exception:
            status = None
        if status is None:
            return {
                "ok": True,
                "kind": "text",
                "text": "Dead Man Switch 未初始化。可在「Dead Man Switch」页设置。",
            }
        return {
            "ok": True,
            "kind": "text",
            "text": f"Dead Man Switch 状态：{getattr(status, 'status', str(status))}",
        }
    except Exception as exc:
        return {
            "ok": True,
            "kind": "text",
            "text": f"Dead Man Switch：未初始化（{exc}）。可在「Dead Man Switch」页设置。",
        }


def _cmd_docs(tokens: list[str]) -> dict[str, Any]:
    """文档管理：/docs list"""
    try:
        from ...doc_extract.extractor import DocumentExtractor

        extractor = DocumentExtractor()
        docs = extractor.list_my_documents(_USER)
        if not docs:
            return {"ok": True, "kind": "text", "text": "暂无文档。可在「文档管理」页上传/提取。"}
        md = ["**已管理文档**\n"]
        for d in docs[:12]:
            di = d.to_dict() if hasattr(d, "to_dict") else d
            md.append(
                f"- **{di.get('title', '?')}**（{di.get('doc_type', di.get('file_type', ''))}）"
            )
        return {"ok": True, "kind": "text", "text": "\n".join(md)}
    except Exception as exc:
        return {"ok": False, "kind": "text", "text": f"文档读取失败: {exc}"}


# =====================================================================
# 执行类命令 + /help 总览（对话傻瓜式操作）
# =====================================================================

_HELP_TEXT = """**对话命令总览**（输入 /命令 即可，无需点页面）

**资源/配置**
- /prompt list | get <名> | set <名> <内容>   —— 提示词
- /expert list | new <id> <名> <人设> | delete <id> —— 自定义专家
- /skill list | enable|disable <名>            —— 技能

**查询/信息**
- /hotline [省份] [功能] · /institution [省] [城市]
- /custom list | get <地区> | presets · /family list | add <姓名>

**业务数据**
- /vault list | add <名称> <类别> · /note list | set <章节> <内容>
- /docs list · /switch status · /task list | add <cron> <内容>

**创作/工具**
- /plot <python代码> · /image <描述> · /browse <网址> · /memorial <姓名> <关系> <回忆>
- /canvas（画布页）· /plot（画图）

**示例**："帮我查北京殡葬热线" → 直接用中文问即可，Agent 会自动调用工具。
"""


async def _cmd_help() -> dict[str, Any]:
    return {"ok": True, "kind": "text", "text": _HELP_TEXT}


async def _cmd_vault_add(tokens: list[str]) -> dict[str, Any]:
    """数字遗产新增：/vault add <名称> <类别(账号/密码/文档/其他)>"""
    if len(tokens) < 2:
        return {"ok": False, "kind": "text", "text": "用法: /vault add <名称> <类别>"}
    name = tokens[0]
    category = tokens[1] if len(tokens) > 1 else "其他"
    try:
        from ...digital_legacy import AssetAction, DigitalAsset, DigitalLegacyStore

        store = DigitalLegacyStore(user_id=_USER)
        store.add_asset(
            DigitalAsset(
                id=f"a_{int(time.time() * 1000) % 10**8}",
                category=category,
                name=name,
                access_hint="",
                action_on_death=AssetAction.DECIDE.value,
            )
        )
        return {
            "ok": True,
            "kind": "text",
            "text": f"已新增数字遗产项：{name}（{category}）。可用 /vault list 查看。",
        }
    except Exception:
        # 兜底走 VaultStore
        from ...vault.store import VaultStore

        try:
            VaultStore().add_item(_USER, category, name, b"", [])
            return {"ok": True, "kind": "text", "text": f"已新增保险库条目：{name}（{category}）。"}
        except Exception as exc2:
            return {"ok": False, "kind": "text", "text": f"新增失败: {exc2}"}


async def _cmd_note_set(tokens: list[str]) -> dict[str, Any]:
    """终活笔记保存：/note set <章节> <内容>"""
    if len(tokens) < 2:
        return {
            "ok": False,
            "kind": "text",
            "text": "用法: /note set <章节(如 遗嘱意愿/医疗意愿/身后安排)> <内容>",
        }
    section = tokens[0]
    content = " ".join(tokens[1:])
    try:
        from datetime import datetime, timezone

        from ...ending_note.models import EndingNote
        from ...ending_note.store import EndingNoteStore

        store = EndingNoteStore()
        note = store.load(_USER)
        if note is None:
            note = EndingNote(
                note_id=f"n-{int(time.time())}",
                user_id=_USER,
                created_at=datetime.now(timezone.utc),
                updated_at=datetime.now(timezone.utc),
            )
        # 写入 personal_info 下的自定义字段（通用文本）
        if not note.personal_info:
            note.personal_info = {}
        note.personal_info[section] = content
        note.updated_at = datetime.now(timezone.utc)
        store.save(note)
        return {"ok": True, "kind": "text", "text": f"已保存终活笔记「{section}」。"}
    except Exception as exc:
        return {"ok": False, "kind": "text", "text": f"保存失败: {exc}"}


async def _cmd_memorial(tokens: list[str]) -> dict[str, Any]:
    """悼文生成：/memorial <姓名> <关系> <回忆…>"""
    if len(tokens) < 2:
        return {
            "ok": False,
            "kind": "text",
            "text": "用法: /memorial <姓名> <关系> <回忆…>，如 /memorial 父亲 儿子 他爱读书、常浇花",
        }
    name = tokens[0]
    relationship = tokens[1]
    memories = "，".join(tokens[2:]).split("、") if len(tokens) > 2 else []
    try:
        from ...memorial_writer.generator import MemorialGenerator
        from ...memorial_writer.models import MemorialRequest

        req = MemorialRequest(
            doc_type="eulogy",
            decedent_name=name,
            relationship=relationship,
            personality_traits=[],
            memories=memories,
            tone="warm",
        )
        result = await MemorialGenerator().generate(req)
        text = getattr(result, "text", "") or "（生成失败）"
        return {"ok": True, "kind": "text", "text": text[:1500]}
    except Exception as exc:
        return {"ok": False, "kind": "text", "text": f"悼文生成失败: {exc}"}
